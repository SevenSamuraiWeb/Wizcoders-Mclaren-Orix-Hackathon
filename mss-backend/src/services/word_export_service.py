"""
Word Document Export Service

Converts credit memo data into formatted Word documents (.docx)
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Dict, Any, List
import logging

logger = logging.getLogger(__name__)


class WordExportService:
    """Service for exporting credit memos to Word format"""

    @staticmethod
    def add_heading(doc: Document, text: str, level: int = 1):
        """Add a styled heading to the document"""
        heading = doc.add_heading(text, level=level)
        return heading

    @staticmethod
    def add_shaded_paragraph(doc: Document, text: str, shade_color: str = "D9E8F5"):
        """Add a paragraph with background shading"""
        paragraph = doc.add_paragraph(text)
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), shade_color)
        paragraph._element.get_or_add_pPr().append(shading_elm)
        return paragraph

    @staticmethod
    def create_credit_memo_word(memo_data: Dict[str, Any], filename: str = "credit_memo.docx") -> bytes:
        """
        Convert credit memo data to a formatted Word document.
        
        Args:
            memo_data: Credit memo dictionary from ReportService.generate_credit_memo()
            filename: Output filename (for reference only, returns bytes)
        
        Returns:
            bytes: Word document file content
        """
        doc = Document()
        
        # Set document margins (1 inch)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # ===== TITLE PAGE =====
        title = doc.add_heading("CREDIT MEMORANDUM", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Company info
        company_para = doc.add_paragraph()
        company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        company_run = company_para.add_run("Financial Analysis Summary\n")
        company_run.font.size = Pt(14)
        company_run.font.bold = True
        
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.add_run(f"Document ID: {memo_data.get('document_id', 'N/A')}")
        
        doc.add_paragraph()  # Spacing
        
        # ===== EXECUTIVE SUMMARY =====
        WordExportService.add_heading(doc, "Executive Summary", level=1)
        
        bullets = memo_data.get("executive_summary_bullets", [])
        for i, bullet in enumerate(bullets, 1):
            if isinstance(bullet, dict):
                # Has confidence tagging
                text = bullet.get("text", str(bullet))
                confidence = bullet.get("source_confidence", "")
                
                para = doc.add_paragraph(style='List Bullet')
                para.add_run(text).bold = True
                
                if confidence:
                    conf_run = para.add_run(f"\n   [Source Confidence: {confidence}]")
                    conf_run.font.italic = True
                    conf_run.font.size = Pt(9)
            else:
                # Plain text
                doc.add_paragraph(text=str(bullet), style='List Bullet')
        
        doc.add_paragraph()  # Spacing
        
        # ===== KEY METRICS TABLE =====
        WordExportService.add_heading(doc, "Key Metrics", level=1)
        
        metrics_table = memo_data.get("metrics_table", {})
        if metrics_table:
            # Create table: 2 columns, 1 header + data rows
            num_rows = len(metrics_table.get("data", [])) + 1
            table = doc.add_table(rows=num_rows, cols=2)
            table.style = 'Light Grid Accent 1'
            
            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = "Metric"
            header_cells[1].text = "Value"
            
            # Make header bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Data rows
            for idx, row_data in enumerate(metrics_table.get("data", []), 1):
                row = table.rows[idx]
                row.cells[0].text = str(row_data.get("metric", ""))
                row.cells[1].text = str(row_data.get("value", ""))
        
        doc.add_paragraph()  # Spacing
        
        # ===== DSCR ANALYSIS =====
        if "key_ratios" in memo_data:
            WordExportService.add_heading(doc, "Debt Service Coverage Ratio Analysis", level=1)
            
            key_ratios = memo_data.get("key_ratios", {})
            dscr_value = key_ratios.get("dscr", {})
            if isinstance(dscr_value, dict):
                dscr = dscr_value.get("value", "N/A")
            else:
                dscr = dscr_value
            
            dscr_para = doc.add_paragraph()
            dscr_para.add_run(f"DSCR: {dscr}").bold = True
            
            interpretation = WordExportService._get_dscr_interpretation(dscr)
            doc.add_paragraph(interpretation)
        
        doc.add_paragraph()  # Spacing
        
        # ===== TOP RISKS =====
        WordExportService.add_heading(doc, "Top 3 Risks", level=1)
        
        risks = memo_data.get("top_3_risks", [])
        for i, risk in enumerate(risks, 1):
            # Risk title with severity
            severity = risk.get("severity", "UNKNOWN")
            risk_para = doc.add_paragraph()
            risk_title = risk_para.add_run(f"{i}. {risk.get('risk_factor', f'Risk {i}')}")
            risk_title.bold = True
            severity_run = risk_para.add_run(f" [{severity}]")
            severity_run.font.color.rgb = WordExportService._get_severity_color(severity)
            
            # Risk description
            description = risk.get("description", "")
            if description:
                doc.add_paragraph(description, style='List Bullet')
            
            # Severity justification
            if "severity_justification" in risk:
                justification = risk.get("severity_justification", "")
                just_para = doc.add_paragraph()
                just_run = just_para.add_run(f"Justification: {justification}")
                just_run.italic = True
                just_run.font.size = Pt(10)
        
        doc.add_paragraph()  # Spacing
        
        # ===== MISSING INFORMATION =====
        if "missing_information" in memo_data:
            missing = memo_data.get("missing_information", [])
            if missing:
                WordExportService.add_heading(doc, "Missing Information", level=1)
                for item in missing:
                    doc.add_paragraph(item, style='List Bullet')
                doc.add_paragraph()
        
        # ===== RED FLAGS =====
        if "red_flags" in memo_data:
            red_flags = memo_data.get("red_flags", [])
            if red_flags:
                WordExportService.add_heading(doc, "Red Flags Detected", level=1)
                for flag in red_flags:
                    flag_para = doc.add_paragraph()
                    flag_title = flag_para.add_run(flag.get("flag", "Flag"))
                    flag_title.bold = True
                    flag_severity = flag.get("severity", "")
                    if flag_severity:
                        sev_run = flag_para.add_run(f" [{flag_severity}]")
                        sev_run.font.color.rgb = WordExportService._get_severity_color(flag_severity)
                    
                    if flag.get("description"):
                        doc.add_paragraph(flag.get("description"), style='List Bullet')
                doc.add_paragraph()
        
        # ===== CREDIT ANALYST CHECKLIST =====
        if "credit_analyst_checklist" in memo_data:
            checklist = memo_data.get("credit_analyst_checklist", {})
            if checklist:
                WordExportService.add_heading(doc, "Credit Analyst Checklist", level=1)
                
                # Reviewed documents
                reviewed = checklist.get("reviewed_documents", {})
                if reviewed:
                    doc.add_paragraph("Documents Reviewed:", style='Heading 3')
                    for doc_name, status in reviewed.items():
                        symbol = "✓" if status else "✗"
                        doc.add_paragraph(f"{symbol} {doc_name}", style='List Bullet')
                
                # Verification steps
                verification = checklist.get("verification_steps", {})
                if verification:
                    doc.add_paragraph("Verification Steps:", style='Heading 3')
                    for step_name, status in verification.items():
                        symbol = "✓" if status else "✗"
                        doc.add_paragraph(f"{symbol} {step_name}", style='List Bullet')
                
                # Overall readiness
                readiness = checklist.get("overall_readiness", "Unknown")
                readiness_para = doc.add_paragraph()
                readiness_para.add_run(f"Overall Readiness: {readiness}").bold = True
                doc.add_paragraph()
        
        # ===== RATIO AVAILABILITY =====
        if "ratio_availability_statement" in memo_data:
            ratio_stmt = memo_data.get("ratio_availability_statement", {})
            if ratio_stmt:
                WordExportService.add_heading(doc, "Ratio Availability Statement", level=1)
                
                # Computable ratios
                computable = ratio_stmt.get("computable_ratios", [])
                if computable:
                    doc.add_paragraph("Computable Ratios:", style='Heading 3')
                    for ratio_info in computable:
                        if isinstance(ratio_info, dict):
                            ratio_name = ratio_info.get("ratio", "")
                            ratio_value = ratio_info.get("value", "")
                            doc.add_paragraph(f"{ratio_name}: {ratio_value}", style='List Bullet')
                        else:
                            doc.add_paragraph(str(ratio_info), style='List Bullet')
                
                # Analysis quality
                quality = ratio_stmt.get("analysis_quality", "")
                if quality:
                    quality_para = doc.add_paragraph()
                    quality_para.add_run(f"Analysis Quality: {quality}").italic = True
                doc.add_paragraph()
        
        # ===== ANALYSIS CONFIDENCE =====
        if "analysis_confidence" in memo_data:
            confidence = memo_data.get("analysis_confidence", {})
            if confidence:
                WordExportService.add_heading(doc, "Analysis Confidence Assessment", level=1)
                
                # Overall confidence
                overall = confidence.get("overall_confidence_level", "Unknown")
                conf_para = doc.add_paragraph()
                conf_para.add_run(f"Overall Confidence Level: {overall}").bold = True
                
                # Completeness score
                completeness = confidence.get("completeness_score", 0)
                doc.add_paragraph(f"Completeness Score: {completeness}/100")
                
                # Data quality
                quality = confidence.get("data_quality_assessment", "")
                if quality:
                    doc.add_paragraph(f"Data Quality: {quality}")
                
                # Missing items
                missing_count = confidence.get("missing_items_count", 0)
                doc.add_paragraph(f"Missing Items: {missing_count}")
                
                # Critical data
                critical = confidence.get("critical_data_available", True)
                critical_text = "Yes" if critical else "No"
                doc.add_paragraph(f"Critical Data Available: {critical_text}")
                doc.add_paragraph()
        
        # ===== DATA SOURCES =====
        if "data_sources" in memo_data:
            sources = memo_data.get("data_sources", {})
            if sources:
                WordExportService.add_heading(doc, "Data Sources & Traceability", level=1)
                
                source_list = sources.get("sources", [])
                for source in source_list:
                    if isinstance(source, dict):
                        source_text = f"{source.get('type', 'Source')}: Page {source.get('page', 'N/A')}"
                    else:
                        source_text = str(source)
                    doc.add_paragraph(source_text, style='List Bullet')
                doc.add_paragraph()
        
        # ===== FOOTER =====
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run("---\nConfidential - For Authorized Use Only")
        footer_run.font.size = Pt(9)
        footer_run.italic = True
        
        # Save to bytes
        from io import BytesIO
        doc_bytes = BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.getvalue()

    @staticmethod
    def _get_dscr_interpretation(dscr: Any) -> str:
        """Get interpretation text for DSCR value"""
        try:
            dscr_float = float(str(dscr).replace("x", ""))
            if dscr_float >= 1.5:
                return "DSCR is strong. The borrower has adequate ability to service debt from operating cash flow with room for earnings decline."
            elif dscr_float >= 1.25:
                return "DSCR is good. The borrower can service debt from operations, but margin for earnings decline is moderate."
            elif dscr_float >= 1.0:
                return "DSCR is adequate. The borrower can cover debt service from operations, but has limited margin for deterioration."
            else:
                return "DSCR is below 1.0. The borrower cannot service debt from operating cash flow alone."
        except:
            return "DSCR analysis not available."

    @staticmethod
    def _get_severity_color(severity: str) -> RGBColor:
        """Get color for severity level"""
        severity = severity.upper()
        if severity == "CRITICAL":
            return RGBColor(192, 0, 0)  # Dark red
        elif severity == "HIGH":
            return RGBColor(255, 0, 0)  # Red
        elif severity == "MEDIUM":
            return RGBColor(255, 165, 0)  # Orange
        elif severity == "LOW":
            return RGBColor(0, 128, 0)  # Green
        else:
            return RGBColor(0, 0, 0)  # Black


__all__ = ["WordExportService"]
