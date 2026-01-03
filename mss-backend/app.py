import streamlit as st
import time
import os
import json
from dotenv import load_dotenv
from groq import Groq
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

# Backend Classes
from pdf_processor import PDFProcessor
from vector_store import PageVectorStore
from retriever import ContextRetriever
from credit_memo_generator import CreditMemoGenerator

# Load environment variables
load_dotenv()


def generate_markdown_summary(memo):
    """Convert memo to markdown format."""
    markdown = f"# Credit Memo Report\n\n"
    
    # Metadata
    if memo.get('metadata'):
        metadata = memo['metadata']
        markdown += f"**Document:** {metadata.get('document_name', 'N/A')}\n"
        markdown += f"**Confidence Level:** {int(metadata.get('overall_confidence', 0) * 100)}%\n"
        markdown += f"**Model:** {metadata.get('model_info', 'N/A')}\n\n"
    
    # Executive Summary
    if memo.get('summary'):
        markdown += "## 1. Executive Summary\n\n"
        summary = memo['summary']
        if summary.get('executive_summary'):
            markdown += f"{summary['executive_summary']}\n\n"
        if summary.get('key_takeaways'):
            markdown += "**Key Takeaways:**\n"
            for item in summary['key_takeaways']:
                markdown += f"- {item}\n"
            markdown += "\n"
        if summary.get('recommendation'):
            markdown += f"**Recommendation:** {summary['recommendation']}\n"
        if summary.get('recommendation_justification'):
            markdown += f"**Justification:** {summary['recommendation_justification']}\n\n"
    
    # Financial Metrics
    if memo.get('financial_metrics'):
        markdown += "## 2. Key Financial Metrics\n\n"
        for metric in memo['financial_metrics']:
            name = f"{metric.get('category', '')} - {metric.get('label', metric.get('name', ''))}"
            value = metric.get('value', 'N/A')
            unit = metric.get('unit', '')
            if unit == '%':
                unit = '%'
            elif unit == 'ratio':
                unit = 'x'
            status = metric.get('status', 'Unknown')
            markdown += f"- **{name}:** {value}{unit} ({status})\n"
        markdown += "\n"
    
    # 5Cs Analysis
    if memo.get('credit_analysis_5cs'):
        markdown += "## 3. Credit Analysis (5Cs)\n\n"
        five_cs = memo['credit_analysis_5cs']
        for key, value in five_cs.items():
            if value:
                markdown += f"### {key.upper()}\n"
                # Get the main content
                content = value.get('assessment') or value.get('equity_position') or value.get('loan_purpose') or value.get('repayment_source') or 'No analysis provided'
                markdown += f"{content}\n\n"
    
    # Risk Assessment
    if memo.get('risk_assessment'):
        markdown += "## 4. Risk Assessment\n\n"
        risks = memo['risk_assessment']
        
        if risks.get('red_flags'):
            markdown += "### Red Flags\n"
            for flag in risks['red_flags']:
                issue = flag.get('issue', '')
                severity = flag.get('severity', 'Medium')
                mitigant = flag.get('mitigant', 'None listed')
                markdown += f"- **{issue}** ({severity} Risk)\n  - Mitigant: {mitigant}\n"
            markdown += "\n"
        
        if risks.get('strengths'):
            markdown += "### Strengths\n"
            for item in risks['strengths']:
                markdown += f"- {item.get('text', '')}\n"
            markdown += "\n"
        
        if risks.get('weaknesses'):
            markdown += "### Weaknesses\n"
            for item in risks['weaknesses']:
                markdown += f"- {item.get('text', '')}\n"
            markdown += "\n"
    
    return markdown


def generate_word_document(memo):
    """Convert memo to Word document format."""
    doc = Document()
    
    # Title
    title = doc.add_heading('Credit Memo Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Metadata
    if memo.get('metadata'):
        metadata = memo['metadata']
        doc.add_paragraph(f"Document: {metadata.get('document_name', 'N/A')}")
        doc.add_paragraph(f"Confidence Level: {int(metadata.get('overall_confidence', 0) * 100)}%")
        doc.add_paragraph(f"Model: {metadata.get('model_info', 'N/A')}")
        doc.add_paragraph()
    
    # Executive Summary
    if memo.get('summary'):
        doc.add_heading('1. Executive Summary', level=1)
        summary = memo['summary']
        
        if summary.get('executive_summary'):
            doc.add_paragraph(summary['executive_summary'])
            doc.add_paragraph()
        
        if summary.get('key_takeaways'):
            doc.add_heading('Key Takeaways', level=2)
            for item in summary['key_takeaways']:
                doc.add_paragraph(item, style='List Bullet')
            doc.add_paragraph()
        
        
        
        if summary.get('recommendation_justification'):
            p = doc.add_paragraph()
            p.add_run('Justification: ').bold = True
            p.add_run(summary['recommendation_justification'])
            doc.add_paragraph()
    
    # Financial Metrics
    if memo.get('financial_metrics'):
        doc.add_heading('2. Key Financial Metrics', level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Category'
        hdr_cells[1].text = 'Metric'
        hdr_cells[2].text = 'Value'
        hdr_cells[3].text = 'Status'
        
        for metric in memo['financial_metrics']:
            row_cells = table.add_row().cells
            row_cells[0].text = metric.get('category', '')
            row_cells[1].text = metric.get('label', metric.get('name', ''))
            value = str(metric.get('value', 'N/A'))
            unit = metric.get('unit', '')
            if unit == '%':
                value += '%'
            elif unit == 'ratio':
                value += 'x'
            row_cells[2].text = value
            row_cells[3].text = metric.get('status', 'Unknown')
        doc.add_paragraph()
    
    # 5Cs Analysis
    if memo.get('credit_analysis_5cs'):
        doc.add_heading('3. Credit Analysis (5Cs)', level=1)
        five_cs = memo['credit_analysis_5cs']
        
        for key, value in five_cs.items():
            if value:
                doc.add_heading(key.upper(), level=2)
                # Get the main content
                content = value.get('assessment') or value.get('equity_position') or value.get('loan_purpose') or value.get('repayment_source') or 'No analysis provided'
                doc.add_paragraph(content)
                doc.add_paragraph()
    
    # Risk Assessment
    if memo.get('risk_assessment'):
        doc.add_heading('4. Risk Assessment', level=1)
        risks = memo['risk_assessment']
        
        if risks.get('red_flags'):
            doc.add_heading('Red Flags', level=2)
            for flag in risks['red_flags']:
                issue = flag.get('issue', '')
                severity = flag.get('severity', 'Medium')
                mitigant = flag.get('mitigant', 'None listed')
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"{issue} ({severity} Risk)\n").bold = True
                p.add_run(f"Mitigant: {mitigant}")
            doc.add_paragraph()
        
        if risks.get('strengths'):
            doc.add_heading('Strengths', level=2)
            for item in risks['strengths']:
                doc.add_paragraph(item.get('text', ''), style='List Bullet')
            doc.add_paragraph()
        
        if risks.get('weaknesses'):
            doc.add_heading('Weaknesses', level=2)
            for item in risks['weaknesses']:
                doc.add_paragraph(item.get('text', ''), style='List Bullet')
            doc.add_paragraph()
    
    if summary.get('recommendation'):
            p = doc.add_paragraph()
            p.add_run('Recommendation: ').bold = True
            p.add_run(summary['recommendation'])
            
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer



def simplify_text(text, api_key):
    """Simplify text using Groq API."""
    try:
        client = Groq(api_key=api_key)
        response = client.chat.completions.create(
            messages=[
                {
                    "role": "system",
                    "content": "You are a professional writer. Simplify the given text while maintaining all important information. Use simpler vocabulary and shorter sentences. Make it easy to understand for non-technical readers."
                },
                {
                    "role": "user",
                    "content": f"Please simplify this text:\n\n{text}"
                }
            ],
            model="meta-llama/llama-4-scout-17b-16e-instruct",
            temperature=0.5,
            max_tokens=2048
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Error simplifying text: {str(e)}"

    # 1. Page Configuration
    st.set_page_config(
        page_title="Auto-Credit Memo Generator",
        page_icon="📄",
        layout="wide",
        initial_sidebar_state="expanded"
    )

    # Custom CSS for "Nice" UI
    st.markdown("""
        <style>
        .main-header {
            font-size: 2.5rem;
            color: #1E3A8A;
            font-weight: 700;
        }
        .section-header {
            font-size: 1.5rem;
            color: #1F2937;
            border-bottom: 2px solid #E5E7EB;
            padding-bottom: 0.5rem;
            margin-top: 1.5rem;
            margin-bottom: 1rem;
        }
        .metric-card {
            background-color: #F3F4F6;
            padding: 1rem;
            border-radius: 0.5rem;
            border-left: 4px solid #3B82F6;
        }
        .risk-card-high {
            background-color: #FEF2F2;
            padding: 1rem;
            border-radius: 0.5rem;
            border: 1px solid #FCA5A5;
            margin-bottom: 0.5rem;
        }
        .risk-card-medium {
            background-color: #FFFBEB;
            padding: 1rem;
            border-radius: 0.5rem;
            border: 1px solid #FCD34D;
            margin-bottom: 0.5rem;
        }
        </style>
    """, unsafe_allow_html=True)

    # 2. Header / Title
    st.markdown('<p class="main-header">📄 AI-Powered Credit Memo Generator</p>', unsafe_allow_html=True)
    st.markdown("Upload a financial statement (PDF) to generate a structured, professional credit memo.")

    # 3. Sidebar for Inputs
    with st.sidebar:
        st.header("Configuration")
        
        # API Key handling
        api_key = os.getenv("GROQ_API_KEY")
        if not api_key:
            api_key = st.text_input("Enter Groq API Key", type="password")
            
        uploaded_file = st.file_uploader("Upload Financial Report (PDF)", type=["pdf"])
        
        st.info("System Status: Ready")
        if uploaded_file:
            st.success("PDF Loaded")

    # 4. Main Area Logic
    if uploaded_file is not None and api_key:
        
        # Initialize session state for generation
        if "memo" not in st.session_state:
            st.session_state.memo = None

        # 5. Generate Button
        if st.button("Generate Credit Memo", type="primary"):
            try:
                with st.spinner("Processing PDF and generating insights... (This may take ~30s)"):
                    # A. Parsing
                    status_text = st.empty()
                    status_text.text("Parsing PDF contents...")
                    bytes_data = uploaded_file.getvalue()
                    processor = PDFProcessor(bytes_data)
                    extracted_data = processor.parse()
                    
                    # B. Vector Indexing
                    status_text.text("Building semantic index...")
                    vector_store = PageVectorStore()
                    vector_store.add_pages(extracted_data)
                    
                    # C. Retrieval & Generation
                    status_text.text("Analyzing sections with AI...")
                    retriever = ContextRetriever(vector_store)
                    client = Groq(api_key=api_key)
                    # Using the specific model as verified
                    generator = CreditMemoGenerator(retriever, client, model_name="meta-llama/llama-4-scout-17b-16e-instruct")
                    
                    # Generate full memo
                    memo = generator.generate_credit_memo()
                    st.session_state.memo = memo
                    
                    status_text.empty()
                    st.toast("Credit Memo Generated Successfully!", icon="✅")
                    
            except Exception as e:
                st.error(f"An error occurred: {str(e)}")

        # 6. Display Results if generated
        if st.session_state.memo:
            memo = st.session_state.memo
            st.divider()
            
            # --- Header Info ---
            c1, c2 = st.columns([3, 1])
            with c1:
                metadata = memo.get('metadata', {})
                st.markdown(f"**Document:** {metadata.get('document_name', 'N/A')}")
            with c2:
                confidence = int(metadata.get('overall_confidence', 0) * 100)
                st.markdown(f"**Confidence:** `{confidence}%`")

            # --- Section 1: Executive Summary ---
            if memo.get('summary'):
                st.markdown('<div class="section-header">1. Executive Summary</div>', unsafe_allow_html=True)
                summary = memo['summary']
                
                if summary.get('executive_summary'):
                    st.markdown(summary['executive_summary'])
                
                if summary.get('key_takeaways'):
                    st.markdown("**Key Takeaways:**")
                    for item in summary['key_takeaways']:
                        st.markdown(f"""
                        <div style="margin-bottom: 8px;">
                            <span style="background-color: #E0F2FE; color: #0369A1; padding: 2px 6px; border-radius: 4px; font-size: 0.8em; font-weight: 600;">•</span>
                            {item}
                        </div>
                        """, unsafe_allow_html=True)
                
                if summary.get('recommendation'):
                    st.markdown(f"**Recommendation:** {summary['recommendation']}")
                if summary.get('recommendation_justification'):
                    st.markdown(f"*{summary['recommendation_justification']}*")

            # --- Section 2: Key Financial Metrics ---
            if memo.get('financial_metrics'):
                st.markdown('<div class="section-header">2. Key Financial Metrics</div>', unsafe_allow_html=True)
                metrics_list = memo['financial_metrics']
                
                # Display in rows of 3
                cols = st.columns(3)
                for i, m in enumerate(metrics_list):
                    with cols[i % 3]:
                        value = str(m.get("value", "N/A"))
                        unit = m.get('unit', '')
                        if unit == '%':
                            value += '%'
                        elif unit == 'ratio':
                            value += 'x'
                        
                        metric_name = f"{m.get('category', '')} - {m.get('label', m.get('name', 'Metric'))}"
                        st.metric(
                            label=metric_name,
                            value=value,
                            delta=f"{m.get('status', 'N/A')}",
                            delta_color="off"
                        )
            
            # --- Section 3: Credit Analysis (5Cs) ---
            if memo.get('credit_analysis_5cs'):
                st.markdown('<div class="section-header">3. Credit Analysis (5Cs)</div>', unsafe_allow_html=True)
                five_cs = memo['credit_analysis_5cs']
                
                for key, value in five_cs.items():
                    if value:
                        st.markdown(f"**{key.upper()}**")
                        content = value.get('assessment') or value.get('equity_position') or value.get('loan_purpose') or value.get('repayment_source') or 'No analysis provided'
                        st.markdown(f"_{content}_")
                        st.markdown("")
            
            # --- Section 4: Risk Assessment ---
            if memo.get('risk_assessment'):
                st.markdown('<div class="section-header">4. Risk Assessment</div>', unsafe_allow_html=True)
                risks = memo['risk_assessment']
                
                if risks.get('red_flags'):
                    st.markdown("**Red Flags:**")
                    for flag in risks['red_flags']:
                        severity = flag.get('severity', 'Medium')
                        issue = flag.get('issue', 'Risk')
                        mitigant = flag.get('mitigant', 'No mitigant listed')
                        
                        css_class = "risk-card-high" if severity == "High" else "risk-card-medium"
                        icon = "🔴" if severity == "High" else "⚠️"
                        
                        st.markdown(f"""
                        <div class="{css_class}">
                            <strong>{icon} {issue}</strong> <span style="color: #6B7280;">({severity} Risk)</span><br>
                            <strong>Mitigant:</strong> {mitigant}
                        </div>
                        """, unsafe_allow_html=True)
                
                col1, col2 = st.columns(2)
                with col1:
                    if risks.get('strengths'):
                        st.markdown("**Strengths:**")
                        for item in risks['strengths']:
                            st.markdown(f"- {item.get('text', '')}")
                
                with col2:
                    if risks.get('weaknesses'):
                        st.markdown("**Weaknesses:**")
                        for item in risks['weaknesses']:
                            st.markdown(f"- {item.get('text', '')}")

            # --- Export and Edit Section ---
            st.divider()
            st.markdown('<div class="section-header">📝 Edit & Export</div>', unsafe_allow_html=True)
            
            # Create columns for buttons and text area
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                if st.button("📄 Download as Markdown", use_container_width=True):
                    md_content = generate_markdown_summary(memo)
                    st.download_button(
                        label="📥 Save Markdown",
                        data=md_content,
                        file_name="credit_memo.md",
                        mime="text/markdown",
                        use_container_width=True
                    )
            
            with col2:
                if st.button("📋 Download as Word", use_container_width=True):
                    doc = generate_word_document(memo)
                    doc_buffer = BytesIO()
                    doc.save(doc_buffer)
                    doc_buffer.seek(0)
                    st.download_button(
                        label="📥 Save Word",
                        data=doc_buffer.getvalue(),
                        file_name="credit_memo.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
            
            with col3:
                if st.button("✨ Simplify Content", use_container_width=True):
                    if "simplified_memo" not in st.session_state:
                        st.session_state.simplified_memo = None
                    
                    with st.spinner("Simplifying content..."):
                        # Simplify the memo content
                        simplified = memo.copy()
                        
                        # Simplify summary section
                        if simplified.get('summary'):
                            if simplified['summary'].get('executive_summary'):
                                simplified['summary']['executive_summary'] = simplify_text(
                                    simplified['summary']['executive_summary'], api_key
                                )
                        
                        st.session_state.simplified_memo = simplified
                        st.success("Content simplified! View below ⬇️")
            
            with col4:
                if st.button("🔄 Reset", use_container_width=True):
                    st.session_state.simplified_memo = None
                    st.rerun()
            
            # Show simplified version if available
            if st.session_state.get("simplified_memo"):
                st.info("ℹ️ Showing simplified version")
                memo = st.session_state.simplified_memo
            
            # Text editing area
            st.markdown("#### ✏️ Edit Summary")
            st.info("You can edit the generated content below and then export it.")
            
            # Create editable text areas for each section
            edited_content = {}
            
            # Executive Summary Editor
            if memo.get('summary'):
                with st.expander("Edit Executive Summary", expanded=False):
                    summary = memo['summary']
                    current_text = summary.get('executive_summary', '')
                    if summary.get('key_takeaways'):
                        current_text += "\\n\\nKey Takeaways:\\n" + "\\n".join([f"- {item}" for item in summary['key_takeaways']])
                    edited_content["executive_summary"] = st.text_area("Executive Summary:", value=current_text, height=150, key="exec_edit")
            
            # Financial Metrics Editor
            if memo.get('financial_metrics'):
                with st.expander("Edit Financial Metrics", expanded=False):
                    current_text = "\\n".join([
                        f"{m.get('category', '')} - {m.get('label', m.get('name', ''))}: {m.get('value', 'N/A')}{m.get('unit', '')} ({m.get('status', '')})"
                        for m in memo['financial_metrics']
                    ])
                    edited_content["financial_metrics"] = st.text_area("Financial Metrics:", value=current_text, height=150, key="metrics_edit")
            
            # 5Cs Editor
            if memo.get('credit_analysis_5cs'):
                with st.expander("Edit 5Cs Analysis", expanded=False):
                    current_text = ""
                    for key, value in memo['credit_analysis_5cs'].items():
                        if value:
                            content = value.get('assessment') or value.get('equity_position') or value.get('loan_purpose') or value.get('repayment_source', '')
                            current_text += f"{key.upper()}:\\n{content}\\n\\n"
                    edited_content["5cs"] = st.text_area("5Cs Analysis:", value=current_text, height=200, key="5cs_edit")
            
            # Risks Editor
            if memo.get('risk_assessment'):
                with st.expander("Edit Risk Assessment", expanded=False):
                    risks = memo['risk_assessment']
                    current_text = ""
                    if risks.get('red_flags'):
                        current_text += "RED FLAGS:\\n"
                        for flag in risks['red_flags']:
                            current_text += f"- {flag.get('issue', '')} ({flag.get('severity', '')}): {flag.get('mitigant', '')}\\n"
                    if risks.get('strengths'):
                        current_text += "\\nSTRENGTHS:\\n"
                        for item in risks['strengths']:
                            current_text += f"- {item.get('text', '')}\\n"
                    if risks.get('weaknesses'):
                        current_text += "\\nWEAKNESSES:\\n"
                        for item in risks['weaknesses']:
                            current_text += f"- {item.get('text', '')}\\n"
                    edited_content["risks"] = st.text_area("Risk Assessment:", value=current_text, height=200, key="risks_edit")
            
            # Export edited content
            st.markdown("#### 📥 Export Edited Content")
            export_col1, export_col2 = st.columns(2)
            
            with export_col1:
                if st.button("Save as Markdown (with edits)", use_container_width=True):
                    edited_md = "# Credit Memo Report (Edited)\\n\\n"
                    metadata = memo.get('metadata', {})
                    edited_md += f"**Document:** {metadata.get('document_name', 'N/A')}\\n\\n"
                    for section_name, content in edited_content.items():
                        if content:
                            edited_md += f"## {section_name.replace('_', ' ').title()}\\n{content}\\n\\n"
                    
                    st.download_button(
                        label="📥 Download Edited Markdown",
                        data=edited_md,
                        file_name="credit_memo_edited.md",
                        mime="text/markdown",
                        use_container_width=True
                    )
            
            with export_col2:
                if st.button("Save as Word (with edits)", use_container_width=True):
                    doc = Document()
                    title = doc.add_heading('Credit Memo Report (Edited)', 0)
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    for section_name, content in edited_content.items():
                        if content:
                            doc.add_heading(section_name.replace('_', ' ').title(), level=1)
                            doc.add_paragraph(content)
                    
                    doc_buffer = BytesIO()
                    doc.save(doc_buffer)
                    doc_buffer.seek(0)
                    
                    st.download_button(
                        label="📥 Download Edited Word",
                        data=doc_buffer.getvalue(),
                        file_name="credit_memo_edited.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )

            # --- Raw Data Expander ---
            with st.expander("View Raw JSON Output"):
                st.json(memo)

    elif not api_key:
        st.warning("Please provide a Groq API Key in the sidebar or .env file to proceed.")
    else:
        st.info("👈 Please upload a PDF document to begin.")

if __name__ == "__main__":
    main()
